# -*- coding: utf-8 -*-
"""
Report Word (.docx) per Scalzamento Locale - Pila da Ponte.
Genera una relazione tecnica professionale in formato DOCX con:
  - Frontespizio con dati di progetto
  - Sezione input (tutti i parametri)
  - Indicatori sintetici
  - Tabella passaggi di calcolo completa
  - Tabella confronto formulazioni
  - Tabella verifiche normative (con colori esito)
  - Analisi di sensitivita' (riferimento ai grafici app)
  - Note tecniche e commenti progettuali
  - Disclaimer e campi di validita'
  - Riferimenti bibliografici
"""
from __future__ import annotations

import datetime
from io import BytesIO
from dataclasses import asdict
from typing import List

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from src import (
    DatiPila, calcola_report, sintesi_indicatori, tabella_passaggi,
    verifiche_scalzamento, commenti_progettuali,
    serie_sensitivita_velocita, serie_sensitivita_tirante,
)


# ============================================================================
# Costanti grafiche
# ============================================================================

COLORE_TITOLO_BG = "2C3E50"        # Grigio scuro heading 1
COLORE_HEAD_BG = "2C3E50"          # Intestazione tabella
COLORE_HEAD_FG = RGBColor(255, 255, 255)
COLORE_VERDE = RGBColor(0x1B, 0x7A, 0x3D)     # OK
COLORE_ROSSO = RGBColor(0xC0, 0x39, 0x2B)     # NON OK
COLORE_GIALLO = RGBColor(0xB8, 0x86, 0x0B)     # ATTENZIONE
COLORE_BLU = RGBColor(0x29, 0x80, 0xB9)        # INFO
COLORE_RIGA_ALT = "F5F6FA"         # Grigio alternato righe


# ============================================================================
# Helper: shading di una cella
# ============================================================================

def _set_cell_shading(cell, color_hex: str) -> None:
    """Imposta il colore di sfondo di una cella tabella."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, bold: bool = False, font_size: int = 8,
                   color: RGBColor | None = None, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Imposta testo con formattazione in una cella."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    # Riduce spazio interno
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)


# ============================================================================
# Sezioni del report
# ============================================================================

def _add_frontespizio(doc: Document, dati: DatiPila) -> None:
    """Sezione 1: frontespizio con titolo, normative e data."""
    # Titolo principale
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RELAZIONE TECNICA")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Scalzamento Locale presso una Pila da Ponte")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # Metadati
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Data: {datetime.date.today().strftime('%d/%m/%Y')}\n"
        f"Software: Scalzamento - Pila da Ponte (CSU/HEC-18, Melville & Coleman 2000)\n"
        f"Normative di riferimento: HEC-18 (5\u00aa ed.), Melville & Coleman (2000), EN 1997"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spaziatore


def _add_input(doc: Document, dati: DatiPila, indicatori: dict) -> None:
    """Sezione 2: parametri di input."""
    h = doc.add_heading("Dati di Input", level=1)
    _style_heading_1(h)

    # Sotto-tabella geometrica
    h2 = doc.add_heading("Geometria della pila e idraulica", level=2)
    parametri = [
        ("Larghezza pila a", f"{dati.larghezza_pila:.4f} m", "Dimensione caratteristica perpendicolare al flusso"),
        ("Lunghezza pila L", f"{dati.lunghezza_pila:.4f} m", "Lunghezza nella direzione del flusso"),
        ("Rapporto L / a", f"{indicatori['L / a [-]']:.4f} \u2013", "Allungamento della pila"),
        ("Forma del naso", dati.forma_naso, "Influenza K1 e il limite round nose"),
        ("Angolo di attacco", f"{dati.angolo_attacco_gradi:.1f}\u00b0", "Angolo tra l'asse del flusso e la pila"),
        ("Tirante a monte y\u2081", f"{dati.tirante_monte:.4f} m", "Profondit\u00e0 idraulica indisturbata"),
        ("Velocit\u00e0 a monte V\u2081", f"{dati.velocita_monte:.4f} m/s", "Velocit\u00e0 media del flusso a monte"),
        ("D50 sedimento", f"{dati.D50_mm:.3f} mm", "Diametro mediano del sedimento di fondo"),
        ("Ss (peso specifico relativo)", f"{dati.Ss:.3f}", "Peso specifico relativo del sedimento"),
    ]
    _add_kv_table(doc, parametri)

    h2 = doc.add_heading("Coefficienti CSU / HEC-18", level=2)
    coeff = [
        ("K1 (forma naso)", f"{dati.k1:.4f}", "Quadrato=1.1, Arrotondato/Cilindro=1.0, Affilato=0.9"),
        ("K2 effettivo", f"{indicatori['K2 effettivo [-]']:.4f}", "Calcolato da angolo e L/a" if dati.usa_k2_automatico else "Inserito manualmente"),
        ("K3 (condizioni fondo)", f"{dati.k3:.4f}", "Plane bed=1.1, Dune medie=1.15, Grandi=1.3"),
        ("K4 (armoring)", f"{dati.k4:.4f}", "1.0 = nessuna riduzione; <1.0 se presente armatura"),
        ("K_total fattoriale", f"{dati.k_total_fattoriale:.4f}", "Coefficiente per formulazione estendibile"),
    ]
    _add_kv_table(doc, coeff)


def _add_indicatori(doc: Document, indicatori: dict) -> None:
    """Sezione 3: indicatori sintetici."""
    h = doc.add_heading("Indicatori Sintetici", level=1)
    _style_heading_1(h)

    items = [
        ("Numero di Froude Fr", f"{indicatori['Fr [-]']:.5f}"),
        ("Rapporto a / y\u2081", f"{indicatori['a / y1 [-]']:.5f}"),
        ("Rapporto L / a", f"{indicatori['L / a [-]']:.5f}"),
        ("K2 effettivo", f"{indicatori['K2 effettivo [-]']:.5f}"),
        ("Velocit\u00e0 critica V\u2091", f"{indicatori['V_c [m/s]']:.5f} m/s"),
        ("Rapporto V / V\u2091", f"{indicatori['V / V_c [-]']:.5f}"),
        ("Regime di trasporto", indicatori["Regime"]),
        ("ys massimo (cautelativo)", f"{indicatori['ys max [m]']:.5f} m"),
        ("ys minimo", f"{indicatori['ys min [m]']:.5f} m"),
        ("Spread (incertezza formulazioni)", f"{indicatori['spread [m]']:.5f} m"),
    ]
    _add_kv_table(doc, items)


def _add_passaggi(doc: Document, df_pass: pd.DataFrame) -> None:
    """Sezione 4: tabella passaggi di calcolo passo per passo."""
    h = doc.add_heading("Passaggi di Calcolo (passo per passo)", level=1)
    _style_heading_1(h)

    note = doc.add_paragraph()
    run = note.add_run(
        "La tabella seguente riporta ogni passaggio intermedio del calcolo "
        "delle profondit\u00e0 di scalzamento, permetendo la verifica manuale dei risultati."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    _add_data_table(doc, df_pass, font_size=7)


def _add_confronto(doc: Document, df_report: pd.DataFrame) -> None:
    """Sezione 5: confronto formulazioni."""
    h = doc.add_heading("Confronto tra Formulazioni", level=1)
    _style_heading_1(h)

    note = doc.add_paragraph()
    run = note.add_run(
        "Confronto dei risultati delle diverse formulazioni di calcolo dello scalzamento locale."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    _add_data_table(doc, df_report, font_size=8)


def _add_verifiche(doc: Document, df_ver: pd.DataFrame) -> None:
    """Sezione 6: verifiche normative con esiti colorati."""
    h = doc.add_heading("Verifiche Normative", level=1)
    _style_heading_1(h)

    # Tabella con colori per esito
    cols = list(df_ver.columns)
    n_cols = len(cols)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, col_name in enumerate(cols):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, COLORE_TITOLO_BG)
        _set_cell_text(cell, col_name, bold=True, font_size=8, color=COLORE_HEAD_FG,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Dati
    for row_idx, (_, row) in enumerate(df_ver.iterrows()):
        table_row = table.add_row()
        esito_val = str(row.get("Esito", ""))

        # Alternanza sfondo
        bg_color = COLORE_RIGA_ALT if row_idx % 2 == 0 else "FFFFFF"

        for i, col_name in enumerate(cols):
            cell = table_row.cells[i]
            val = row[col_name]
            txt = f"{val:.4f}" if isinstance(val, float) else str(val)

            # Determina colore esito
            font_color = None
            if col_name == "Esito":
                if esito_val == "OK":
                    font_color = COLORE_VERDE
                elif esito_val == "NON OK":
                    font_color = COLORE_ROSSO
                elif esito_val == "ATTENZIONE":
                    font_color = COLORE_GIALLO
                elif esito_val == "INFO":
                    font_color = COLORE_BLU

            _set_cell_text(cell, txt, bold=(col_name == "Esito"),
                          font_size=7, color=font_color,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER if col_name in ("N.", "Esito") else WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_shading(cell, bg_color)


def _add_sensibilita(doc: Document, dati: DatiPila) -> None:
    """Sezione 7: analisi di sensitivita'."""
    h = doc.add_heading("Analisi di Sensibilit\u00e0", level=1)
    _style_heading_1(h)

    df_sens_v = serie_sensitivita_velocita(dati)
    df_sens_y = serie_sensitivita_tirante(dati)

    p = doc.add_paragraph()
    run = p.add_run(
        "L'analisi di sensitivit\u00e0 esplora la variazione della profondit\u00e0 "
        "di scalzamento ys al variare della velocit\u00e0 V\u2081 e del tirante y\u2081. "
        "I grafici interattivi sono disponibili nell'applicazione Streamlit (tab Grafici)."
    )
    run.font.size = Pt(9)

    h2 = doc.add_heading("Sensibilit\u00e0 rispetto alla velocit\u00e0", level=2)

    # Riepilogo sintetico velocit\u00e0
    for formula_name in df_sens_v["Formulazione"].unique():
        sub = df_sens_v[df_sens_v["Formulazione"] == formula_name]
        ys_min = sub["ys [m]"].min()
        ys_max = sub["ys [m]"].max()
        p = doc.add_paragraph()
        run = p.add_run(
            f"  {formula_name}: ys da {ys_min:.4f} m a {ys_max:.4f} m "
            f"(range {ys_max - ys_min:.4f} m)"
        )
        run.font.size = Pt(9)

    h2 = doc.add_heading("Sensibilit\u00e0 rispetto al tirante", level=2)

    for formula_name in df_sens_y["Formulazione"].unique():
        sub = df_sens_y[df_sens_y["Formulazione"] == formula_name]
        ys_min = sub["ys [m]"].min()
        ys_max = sub["ys [m]"].max()
        p = doc.add_paragraph()
        run = p.add_run(
            f"  {formula_name}: ys da {ys_min:.4f} m a {ys_max:.4f} m "
            f"(range {ys_max - ys_min:.4f} m)"
        )
        run.font.size = Pt(9)

    # Tabelle dettagliate (punti campionati)
    h2 = doc.add_heading("Dettaglio sensitivit\u00e0 velocit\u00e0", level=2)
    _add_data_table(doc, df_sens_v, font_size=7)

    doc.add_page_break()
    h2 = doc.add_heading("Dettaglio sensitivit\u00e0 tirante", level=2)
    _add_data_table(doc, df_sens_y, font_size=7)


def _add_note_tecniche(doc: Document, note: List[str]) -> None:
    """Sezione 8: note tecniche e commenti progettuali."""
    h = doc.add_heading("Note Tecniche e Commenti Progettuali", level=1)
    _style_heading_1(h)

    for i, item in enumerate(note, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {item}")
        run.font.size = Pt(9)


def _add_disclaimer(doc: Document) -> None:
    """Sezione 9: disclaimer e campi di validita'."""
    h = doc.add_heading("Disclaimer e Campi di Validit\u00e0", level=1)
    _style_heading_1(h)

    disclaimer_items = [
        "I risultati prodotti dall'applicazione devono essere verificati da un ingegnere abilitato "
        "all'esercizio della professione e non sostituiscono in alcun modo il giudizio professionale "
        "del progettista responsabile.",

        "Le formule CSU/HEC-18 e Melville & Coleman (2000) sono state sviluppate per condizioni "
        "subcritiche (Fr < 1) e per pile singole in alvei rettilinei a bassa sinuosita'. "
        "L'applicabilita' in condizioni supercritiche, alvei meandriformi, o per gruppi di pile "
        "dev'essere valutata caso per caso con letteratura specifica.",

        "La formula CSU/HEC-18 applica un limite massimo per pile a naso arrotondato allineate "
        "al flusso: ys <= 2.4a per Fr <= 0.8, ys <= 3.0a per Fr > 0.8.",

        "I coefficienti K1\u2026K4 sono definiti per condizioni specifiche; valori al di fuori "
        "degli intervalli raccomandati da HEC-18 richiedono giustificazione idraulica.",

        "La condizione clear-water (V <= V\u2091) produce scalzamento progressivo verso un valore "
        "di equilibrio. La condizione live-bed (V > V\u2091) produce scalzamento dinamico con "
        "riempimento parziale durante il decremento delle portate.",

        "Lo scalzamento conservativo da usare in progetto e' il valore massimo tra tutte le "
        "formulazioni, applicando un adeguato margine di sicurezza.",

        "Il presente software non tiene conto di effetti locali quali deposizione a monte, "
        "erosione generale dell'alveo, o interazione con altri ostacoli idraulici.",
    ]

    for item in disclaimer_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _add_riferimenti(doc: Document) -> None:
    """Sezione 10: riferimenti bibliografici."""
    h = doc.add_heading("Riferimenti Bibliografici", level=1)
    _style_heading_1(h)

    rif_item = [
        "[1] Richardson, E.V., Davis, S.R. (2001). Evaluating Scour at Bridges, "
        "Hydraulic Engineering Circular No. 18 (HEC-18), 4th ed., FHWA-IP-90-017, "
        "Federal Highway Administration, U.S. Department of Transportation.",

        "[2] Arneson, L.A., Zevenbergen, L.W., Lagasse, P.F., Clopper, P.E. (2012). "
        "Evaluating Scour at Bridges, HEC-18, 5th ed., FHWA-HIF-12-003, "
        "Federal Highway Administration.",

        "[3] Melville, B.W., Coleman, S.E. (2000). Bridge Scour, Water Resources "
        "Publications, LLC, Highlands Ranch, Colorado, USA.",

        "[4] EN 1997-1:2004 (Eurocodice 7). Progettazione geotecnica - Parte 1: "
        "Regole generali. CEN, Bruxelles.",

        "[5] AASHTO (2020). LRFD Bridge Design Specifications, 9th ed., "
        "American Association of State Highway and Transportation Officials.",

        "[6] Raudkivi, A.J. (1998). Loose Boundary Hydraulics, A.A. Balkema, "
        "Rotterdam, 4th ed.",

        "[7] Breusers, H.N.C., Nicollet, G., Shen, H.W. (1977). Local scour around "
        "cylindrical piers. Journal of Hydraulic Research, 15(3), 211-252.",
    ]

    for item in rif_item:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ============================================================================
# Helper: tabelle
# ============================================================================

def _add_kv_table(doc: Document, items: list) -> None:
    """Aggiunge una tabella chiave-valore-descrizione (Parametro | Valore | Riferimento)."""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Imposta larghezze colonne
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(4.0)
        row.cells[2].width = Cm(7.5)

    # Header
    header = table.rows[0]
    for i, label in enumerate(["Parametro", "Valore", "Riferimento / Descrizione"]):
        cell = header.cells[i]
        _set_cell_shading(cell, COLORE_TITOLO_BG)
        _set_cell_text(cell, label, bold=True, font_size=8, color=COLORE_HEAD_FG,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Righe dati
    for idx, item in enumerate(items):
        if len(item) == 3:
            key, value, ref = item
        else:
            key, value = item
            ref = ""
        row = table.add_row()
        bg = COLORE_RIGA_ALT if idx % 2 == 0 else "FFFFFF"
        _set_cell_text(row.cells[0], key, bold=True, font_size=8)
        _set_cell_text(row.cells[1], value, font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[2], ref, font_size=7)
        for c in range(3):
            _set_cell_shading(row.cells[c], bg)

    doc.add_paragraph()  # spaziatore


def _add_data_table(doc: Document, df: pd.DataFrame, font_size: int = 8) -> None:
    """Aggiunge una tabella generica da DataFrame con header colorato e righe alternate."""
    cols = list(df.columns)
    n_cols = len(cols)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, col_name in enumerate(cols):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, COLORE_TITOLO_BG)
        _set_cell_text(cell, col_name, bold=True, font_size=font_size, color=COLORE_HEAD_FG,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Dati
    for row_idx, (_, row) in enumerate(df.iterrows()):
        table_row = table.add_row()
        bg = COLORE_RIGA_ALT if row_idx % 2 == 0 else "FFFFFF"
        for i, col_name in enumerate(cols):
            val = row[col_name]
            txt = f"{val:.4f}" if isinstance(val, float) else str(val)
            _set_cell_text(table_row.cells[i], txt, font_size=font_size,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_shading(table_row.cells[i], bg)

    doc.add_paragraph()  # spaziatore


# ============================================================================
# Helper: heading styling
# ============================================================================

def _style_heading_1(heading) -> None:
    """Applica stile grigio scuro con testo bianco al livello 1."""
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        run.font.size = Pt(14)
        run.bold = True
    # Aggiunge sfondo tramite shading del paragrafo
    pPr = heading._element.get_or_add_pPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{COLORE_TITOLO_BG}" w:val="clear"/>'
    )
    pPr.append(shading)


# ============================================================================
# Funzione principale
# ============================================================================

def create_enhanced_word_report(dati: DatiPila) -> bytes:
    """
    Genera un report Word (.docx) professionale per la verifica dello
    scalzamento locale. Ritorna i byte del file DOCX in memoria.
    """
    doc = Document()

    # Stile globale
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    # Raccolta dati
    indicatori = sintesi_indicatori(dati)
    df_pass = tabella_passaggi(dati)
    df_report = calcola_report(dati)
    df_ver = verifiche_scalzamento(dati)
    note = commenti_progettuali(dati)

    # 1. Frontespizio
    _add_frontespizio(doc, dati)
    doc.add_page_break()

    # 2. Dati di input
    _add_input(doc, dati, indicatori)

    # 3. Indicatori sintetici
    _add_indicatori(doc, indicatori)

    # 4. Passaggi di calcolo
    doc.add_page_break()
    _add_passaggi(doc, df_pass)

    # 5. Confronto formulazioni
    doc.add_page_break()
    _add_confronto(doc, df_report)

    # 6. Verifiche normative
    doc.add_page_break()
    _add_verifiche(doc, df_ver)

    # 7. Analisi di sensitivita'
    doc.add_page_break()
    _add_sensibilita(doc, dati)

    # 8. Note tecniche
    _add_note_tecniche(doc, note)

    # 9. Disclaimer
    doc.add_page_break()
    _add_disclaimer(doc)

    # 10. Riferimenti bibliografici
    _add_riferimenti(doc)

    # Esportazione in memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()